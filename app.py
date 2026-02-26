import streamlit as st
import os
import time
import PyPDF2
import json
import pandas as pd
import plotly.express as px
from io import BytesIO
from docx import Document
from google import genai
import gspread
from google.oauth2.service_account import Credentials
from google.cloud import storage

st.set_page_config(page_title="RBS Knowledge Corner", layout="wide", page_icon="🏛️")

# --- Authentication & Cloud Setup ---
GOOGLE_CREDS_JSON = os.environ.get("GOOGLE_CREDENTIALS_JSON")
SHEET_ID = os.environ.get("SPREADSHEET_ID")
GCS_BUCKET_NAME = os.environ.get("GCS_BUCKET_NAME")
api_key = os.environ.get("GEMINI_API_KEY")

@st.cache_resource
def get_google_clients():
    if not GOOGLE_CREDS_JSON or not SHEET_ID or not GCS_BUCKET_NAME:
        return None, None
    try:
        creds_dict = json.loads(GOOGLE_CREDS_JSON)
        scopes = ['https://www.googleapis.com/auth/spreadsheets']
        sheet_creds = Credentials.from_service_account_info(creds_dict, scopes=scopes)
        gc = gspread.authorize(sheet_creds)
        sh = gc.open_by_key(SHEET_ID)
        storage_client = storage.Client.from_service_account_info(creds_dict)
        return sh, storage_client
    except Exception as e:
        st.error(f"Google Auth Error: {e}")
        return None, None

sh, storage_client = get_google_clients()

def init_sheets():
    if sh:
        try:
            worksheet_titles = [ws.title for ws in sh.worksheets()]
            
            if "Judgments" not in worksheet_titles:
                sh.add_worksheet(title="Judgments", rows="1000", cols="10")
            judgments_sheet = sh.worksheet("Judgments")
            if not judgments_sheet.row_values(1):
                judgments_sheet.append_row(["ID", "Case Name", "Act Name", "Section Number", "Authority", "Brief Facts", "Decision Held", "PDF File IDs", "AI Notes", "Status"])
                
            if "Internal Usage" not in worksheet_titles:
                sh.add_worksheet(title="Internal Usage", rows="1000", cols="10")
            internal_sheet = sh.worksheet("Internal Usage")
            if not internal_sheet.row_values(1):
                internal_sheet.append_row(["ID", "Judgment ID", "Internal Matter Name", "Internal Notice", "Usage Notes", "AI Brief"])
                
            if "Notice Replies" not in worksheet_titles:
                sh.add_worksheet(title="Notice Replies", rows="1000", cols="10")
            notice_sheet = sh.worksheet("Notice Replies")
            if not notice_sheet.row_values(1):
                notice_sheet.append_row(["ID", "Matter Name", "Notice Text", "Internal Judgments Used", "External References", "Final Reply"])
                
        except Exception as e:
            st.error(f"Error initializing sheets: {e}")

if sh:
    init_sheets()

# --- Google Cloud Storage File Handlers ---
def upload_to_gcs(file_buffer, file_name):
    try:
        bucket = storage_client.bucket(GCS_BUCKET_NAME)
        blob = bucket.blob(file_name)
        file_buffer.seek(0)
        blob.upload_from_file(file_buffer, content_type='application/pdf')
        return file_name
    except Exception as e:
        st.error(f"GCS Upload Error: {e}")
        return None

def download_from_gcs(file_name):
    try:
        bucket = storage_client.bucket(GCS_BUCKET_NAME)
        blob = bucket.blob(file_name)
        return blob.download_as_bytes()
    except Exception as e:
        return None

# --- Helper Functions ---
def extract_text_from_buffers(pdf_buffers):
    text = ""
    for pdf_buffer in pdf_buffers:
        try:
            pdf_buffer.seek(0)
            reader = PyPDF2.PdfReader(pdf_buffer)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        except Exception:
            pass
    return text

def ask_ai(prompt):
    if not api_key:
        return None, "Error: API Key is missing from Environment Variables."
    try:
        client = genai.Client(api_key=api_key)
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt
        )
        return response.text, None
    except Exception as e:
        return None, f"AI Error: {e}"

def create_word_docx(text, title="Legal Document"):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(text)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- Initialize Session State ---
if 'form_data' not in st.session_state:
    st.session_state.form_data = {
        "case_name": "", "act_name": "", "section_number": "", 
        "authority": "", "brief_facts": "", "decision_held": "", "ai_notes": ""
    }
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'notice_text' not in st.session_state:
    st.session_state.notice_text = ""
if 'suggested_cases' not in st.session_state:
    st.session_state.suggested_cases = []
if 'drafted_reply' not in st.session_state:
    st.session_state.drafted_reply = ""

# --- UI Layout ---
if not sh or not storage_client:
    st.error("🚨 System not connected to Google Cloud. Please configure Environment Variables on Render.")
    st.stop()

st.title("🏛️ RBS Knowledge Corner")

tab_dash, tab_search, tab_matters, tab_add, tab_reply, tab_chat = st.tabs([
    "📊 Dashboard", "🔍 Search & Edit", "📂 Internal Matters", "➕ Add Judgment", "📝 Draft Reply", "💬 Chat with PDF"
])

# ==========================================
# TAB 1: DASHBOARD
# ==========================================
with tab_dash:
    st.header("Firm Analytics")
    try:
        j_data = sh.worksheet("Judgments").get_all_records()
        i_data = sh.worksheet("Internal Usage").get_all_records()
        r_data = sh.worksheet("Notice Replies").get_all_records()
        df_j = pd.DataFrame(j_data)
        
        col1, col2, col3 = st.columns(3)
        col1.metric("Total Judgments Banked", len(j_data))
        col2.metric("Quick Links", len(i_data))
        col3.metric("Drafted Notice Replies", len(r_data))
        
        if not df_j.empty:
            st.markdown("---")
            c1, c2, c3 = st.columns(3)
            with c1:
                if 'Act Name' in df_j.columns and not df_j['Act Name'].replace('', pd.NA).dropna().empty:
                    fig1 = px.pie(df_j[df_j['Act Name'] != ''], names='Act Name', title='Judgments by Act', hole=0.4)
                    st.plotly_chart(fig1, use_container_width=True)
            with c2:
                if 'Authority' in df_j.columns and not df_j['Authority'].replace('', pd.NA).dropna().empty:
                    fig2 = px.histogram(df_j[df_j['Authority'] != ''], x='Authority', title='Judgments by Authority')
                    st.plotly_chart(fig2, use_container_width=True)
            with c3:
                if 'Status' in df_j.columns:
                    fig3 = px.pie(df_j, names='Status', title='Law Status Distribution')
                    st.plotly_chart(fig3, use_container_width=True)
    except Exception as e:
        st.info("Dashboard will populate as data is added.")

# ==========================================
# TAB 2: SEARCH, EDIT & LOGS
# ==========================================
with tab_search:
    st.header("Search, Edit, and Review Logs")
    search_term = st.text_input("Universal Search (Case Name, Facts, Decision):").lower()
    try:
        judgments_sheet = sh.worksheet("Judgments")
        judgments = judgments_sheet.get_all_records()
        internal_uses = sh.worksheet("Internal Usage").get_all_records()
        replies = sh.worksheet("Notice Replies").get_all_records()
        
        results = []
        if search_term:
            for row in judgments:
                if (search_term in str(row.get('Case Name', '')).lower() or 
                    search_term in str(row.get('Brief Facts', '')).lower() or 
                    search_term in str(row.get('Decision Held', '')).lower()):
                    results.append(row)
        else:
            results = judgments
            
        if results:
            st.success(f"Showing {len(results)} judgment(s).")
            for row in results:
                j_id = row.get("ID")
                c_name = row.get("Case Name")
                status = row.get("Status", "🟢 Good Law")
                
                with st.expander(f"{status} | {c_name} | {row.get('Act Name')} - Sec {row.get('Section Number')}"):
                    if "🛑" in status:
                        st.error("WARNING: This judgment has been marked as Overruled or Bad Law.")
                    
                    st.markdown(f"**Authority:** {row.get('Authority')}")
                    st.markdown(f"**Brief Facts:**\n{row.get('Brief Facts')}")
                    st.markdown(f"**Decision Held:**\n{row.get('Decision Held')}")
                    
                    # INTERNAL USAGE LOG
                    st.markdown("---")
                    st.markdown("### 📋 Internal Usage Log")
                    use_count = 0
                    
                    # Check Quick Links
                    for use in internal_uses:
                        if str(use.get('Judgment ID')) == str(j_id):
                            use_count += 1
                            st.markdown(f"- **Linked Matter:** {use.get('Internal Matter Name')} | *Notes: {use.get('Usage Notes')}*")
                    
                    # Check Replies
                    for rep in replies:
                        if c_name in str(rep.get('Internal Judgments Used', '')):
                            use_count += 1
                            st.markdown(f"- **Drafted Reply For:** {rep.get('Matter Name')}")
                            
                    if use_count == 0:
                        st.caption("This judgment has not been cited in any internal matters yet.")
                    
                    # EDIT FUNCTIONALITY
                    st.markdown("---")
                    edit_mode = st.checkbox(f"✏️ Edit '{c_name}'", key=f"edit_check_{j_id}")
                    if edit_mode:
                        with st.form(f"edit_form_{j_id}"):
                            e_c_name = st.text_input("Case Name", value=c_name)
                            e_act = st.text_input("Act Name", value=row.get('Act Name'))
                            e_sec = st.text_input("Section", value=row.get('Section Number'))
                            e_auth = st.text_input("Authority", value=row.get('Authority'))
                            e_status = st.selectbox("Status", ["🟢 Good Law", "🟡 Distinguished / Caution", "🛑 Overruled / Bad Law"], index=["🟢 Good Law", "🟡 Distinguished / Caution", "🛑 Overruled / Bad Law"].index(status) if status in ["🟢 Good Law", "🟡 Distinguished / Caution", "🛑 Overruled / Bad Law"] else 0)
                            e_facts = st.text_area("Brief Facts", value=row.get('Brief Facts'))
                            e_decision = st.text_area("Decision Held", value=row.get('Decision Held'))
                            
                            if st.form_submit_button("💾 Save Changes"):
                                try:
                                    cell = judgments_sheet.find(str(j_id))
                                    # Update cells in the specific row (Columns B through J, assuming A is ID)
                                    judgments_sheet.update(f"B{cell.row}:J{cell.row}", [[e_c_name, e_act, e_sec, e_auth, e_facts, e_decision, row.get('PDF File IDs'), row.get('AI Notes'), e_status]])
                                    st.success("Judgment updated successfully! Please refresh to see changes.")
                                except Exception as e:
                                    st.error(f"Error updating sheet: {e}")

                    # Attachments
                    file_ids = str(row.get("PDF File IDs", "")).split(",")
                    if file_ids and file_ids[0] != "":
                        st.markdown("**Attachments:**")
                        for idx, fid in enumerate(file_ids):
                            if fid.strip():
                                file_bytes = download_from_gcs(fid.strip())
                                if file_bytes:
                                    st.download_button(label=f"⬇️ Download PDF {idx+1}", data=file_bytes, file_name=f"{c_name}_Part{idx+1}.pdf", mime="application/pdf", key=f"dl_{fid}")
    except Exception as e:
        st.warning("Could not fetch records.")

# ==========================================
# TAB 3: INTERNAL MATTERS DASHBOARD
# ==========================================
with tab_matters:
    st.header("📂 Internal Client Matters & Submissions")
    st.markdown("Review all active and historical matters your firm has logged.")
    
    try:
        replies_data = sh.worksheet("Notice Replies").get_all_records()
        links_data = sh.worksheet("Internal Usage").get_all_records()
        
        # Combine unique matter names
        all_matters = list(set([r.get('Matter Name') for r in replies_data if r.get('Matter Name')] + 
                               [l.get('Internal Matter Name') for l in links_data if l.get('Internal Matter Name')]))
        
        if all_matters:
            selected_matter = st.selectbox("Select a Matter / Client to Review:", ["-- Select --"] + sorted(all_matters))
            
            if selected_matter != "-- Select --":
                st.markdown(f"### Matter: {selected_matter}")
                
                # Show Drafted Replies for this matter
                matter_replies = [r for r in replies_data if r.get('Matter Name') == selected_matter]
                if matter_replies:
                    st.subheader("📝 Notice Replies & Submissions")
                    for rep in matter_replies:
                        with st.expander(f"Reply Drafted (ID: {rep.get('ID')})"):
                            st.markdown("**Original Notice Received:**")
                            st.caption(rep.get('Notice Text')[:500] + "...")
                            st.markdown(f"**RBS Precedents Used:** {rep.get('Internal Judgments Used')}")
                            st.markdown(f"**External References:** {rep.get('External References')}")
                            st.markdown("**Final Reply:**")
                            st.info(rep.get('Final Reply'))
                            
                            docx_file = create_word_docx(rep.get('Final Reply'), f"Reply - {selected_matter}")
                            st.download_button("📄 Download Reply as Word", data=docx_file, file_name=f"Reply_{selected_matter}.docx", key=f"dl_rep_{rep.get('ID')}")
                
                # Show Quick Links for this matter
                matter_links = [l for l in links_data if l.get('Internal Matter Name') == selected_matter]
                if matter_links:
                    st.subheader("🔗 Linked Research & Precedents")
                    for link in matter_links:
                        st.markdown(f"- **Judgment ID Cited:** {link.get('Judgment ID')}")
                        st.markdown(f"  - *Strategy/Notes:* {link.get('Usage Notes')}")
        else:
            st.info("No internal matters have been logged yet. Use the 'Draft Reply' tab to start.")
            
    except Exception as e:
        st.error(f"Error loading internal matters: {e}")

# ==========================================
# TAB 4: ADD NEW JUDGMENT
# ==========================================
with tab_add:
    st.header("1. Upload & AI Auto-Fill")
    uploaded_files = st.file_uploader("Upload Judgments (PDF)", type=["pdf"], accept_multiple_files=True)
    
    if st.button("🤖 AI: Read PDFs & Auto-Fill"):
        if uploaded_files:
            with st.spinner("Extracting details..."):
                pdf_text = extract_text_from_buffers(uploaded_files)
                prompt = f"""Extract details from this judgment into JSON format ONLY with keys: "case_name", "act_name", "section_number", "authority", "brief_facts", "decision_held", "ai_notes". Text: {pdf_text[:30000]}"""
                res, err = ask_ai(prompt)
                if not err:
                    try:
                        data = json.loads(res.replace("```json", "").replace("```", "").strip())
                        for key in st.session_state.form_data.keys():
                            st.session_state.form_data[key] = data.get(key, "")
                        st.success("Auto-filled below!")
                    except:
                        st.error("AI output formatting failed. Fill manually.")
        else:
            st.warning("Upload files first.")

    st.header("2. Review & Save to Cloud")
    with st.form("add_judgment_form", clear_on_submit=False):
        c1, c2, c3 = st.columns(3)
        with c1:
            case_name = st.text_input("Name of Case *", value=st.session_state.form_data["case_name"])
            act_name = st.text_input("Act Name", value=st.session_state.form_data["act_name"])
        with c2:
            section_num = st.text_input("Section Number", value=st.session_state.form_data["section_number"])
            authority = st.text_input("Authority", value=st.session_state.form_data["authority"])
        with c3:
            status = st.selectbox("Current Status", ["🟢 Good Law", "🟡 Distinguished / Caution", "🛑 Overruled / Bad Law"])
            
        brief_facts = st.text_area("Brief Facts *", value=st.session_state.form_data["brief_facts"])
        decision_held = st.text_area("Decision Held *", value=st.session_state.form_data["decision_held"])
        ai_notes = st.text_area("AI Notes", value=st.session_state.form_data["ai_notes"])
        
        if st.form_submit_button("✅ Upload & Save to Cloud"):
            if case_name and brief_facts and decision_held:
                with st.spinner("Uploading to Google Cloud Storage and saving to Sheets..."):
                    j_id = str(int(time.time()))
                    gcs_file_ids = []
                    if uploaded_files:
                        for f in uploaded_files:
                            file_buffer = BytesIO(f.getbuffer())
                            fid = upload_to_gcs(file_buffer, f"{j_id}_{f.name}")
                            if fid: gcs_file_ids.append(fid)
                    
                    row_data = [j_id, case_name, act_name, section_num, authority, brief_facts, decision_held, ",".join(gcs_file_ids), ai_notes, status]
                    sh.worksheet("Judgments").append_row(row_data)
                    st.session_state.form_data = {k: "" for k in st.session_state.form_data}
                    st.success("Saved successfully to the Cloud!")

# ==========================================
# TAB 5: DRAFT NOTICE REPLY
# ==========================================
with tab_reply:
    st.header("📝 Step 1: Analyze Notice")
    notice_files = st.file_uploader("Upload Legal Notice(s) received (PDF)", type=["pdf"], accept_multiple_files=True, key="notice_uploader")
    
    if st.button("🔍 Read Notice & Suggest Strategies"):
        if notice_files:
            with st.spinner("Reading Notice and searching RBS Knowledge Corner..."):
                st.session_state.notice_text = extract_text_from_buffers(notice_files)
                
                all_judgments = sh.worksheet("Judgments").get_all_records()
                good_law_catalog = ""
                for j in all_judgments:
                    if "Good Law" in j.get("Status", ""):
                        good_law_catalog += f"ID: {j['ID']} | Case: {j['Case Name']} | Facts: {j['Brief Facts']} | Decision: {j['Decision Held']}\n\n"
                
                prompt = f"""
                You are a senior litigation attorney. 
                Read this legal notice: {st.session_state.notice_text[:15000]}
                
                Task 1: Identify the best internal precedents from this catalog:
                {good_law_catalog[:30000]}
                
                Task 2: Suggest 2 or 3 major EXTERNAL landmark legal precedents (not in the catalog) that are highly relevant to defending against this notice.
                
                Return ONLY a valid JSON object with two keys:
                "internal_cases": [list of exact Case Names from the catalog]
                "external_suggestions": [list of external case names and a 1-sentence explanation of why]
                """
                res, err = ask_ai(prompt)
                
                if not err:
                    try:
                        suggestions = json.loads(res.replace("```json", "").replace("```", "").strip())
                        st.session_state.suggested_cases = suggestions.get("internal_cases", [])
                        st.success("Analysis complete!")
                        if suggestions.get("external_suggestions"):
                            st.info("**AI Suggested External Precedents:**\n" + "\n".join(suggestions.get("external_suggestions", [])))
                    except:
                        st.warning("AI provided suggestions, but formatting was off. Please select manually below.")
        else:
            st.warning("Please upload a notice first.")

    st.markdown("---")
    st.header("📝 Step 2: Build Your Argument")
    
    try:
        all_judgments = sh.worksheet("Judgments").get_all_records()
        all_case_names = [j['Case Name'] for j in all_judgments]
        default_selections = [c for c in st.session_state.suggested_cases if c in all_case_names]
        
        selected_internal = st.multiselect("Select RBS Precedents to include:", options=all_case_names, default=default_selections)
        
        st.markdown("**Add External References (From AI suggestions or your own knowledge):**")
        external_refs = st.text_area("Type any outside case laws or specific arguments you want included in the draft:")
        
        if st.button("✍️ Draft Reply"):
            if st.session_state.notice_text:
                with st.spinner("Drafting your response..."):
                    selected_details = ""
                    for j in all_judgments:
                        if j['Case Name'] in selected_internal:
                            selected_details += f"Case: {j['Case Name']}\nRuling: {j['Decision Held']}\n\n"
                    
                    draft_prompt = f"""
                    You are an expert legal counsel. Draft a formal, professional legal reply to the following notice.
                    
                    Original Notice received:
                    {st.session_state.notice_text[:15000]}
                    
                    You MUST cite and apply these internal precedents to support our position:
                    {selected_details}
                    
                    You MUST ALSO incorporate these specific external precedents/notes:
                    {external_refs}
                    
                    Draft the full body of the legal reply. Use standard legal formatting and authoritative tone. Do not use placeholders for dates/names if you can deduce them.
                    """
                    
                    draft_res, err = ask_ai(draft_prompt)
                    if not err:
                        st.session_state.drafted_reply = draft_res
            else:
                st.error("Please upload and analyze a Notice in Step 1 first.")

    except Exception as e:
        st.error("Error loading precedents.")

    st.markdown("---")
    st.header("📝 Step 3: Review, Edit, and Save")
    
    matter_name = st.text_input("Matter / Client Name (For tracking):")
    final_draft = st.text_area("Edit your Final Reply:", value=st.session_state.drafted_reply, height=400)
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("💾 Save to RBS Knowledge Corner"):
            if matter_name and final_draft:
                with st.spinner("Saving to Google Sheets..."):
                    record_id = str(int(time.time()))
                    row_data = [
                        record_id, 
                        matter_name, 
                        st.session_state.notice_text, 
                        ", ".join(selected_internal), 
                        external_refs, 
                        final_draft
                    ]
                    sh.worksheet("Notice Replies").append_row(row_data)
                    st.success("Notice and Reply successfully recorded! You can view it in the 'Internal Matters' tab.")
            else:
                st.error("Please provide a Matter Name and ensure the draft is not empty.")
    
    with col2:
        if final_draft:
            docx_file = create_word_docx(final_draft, f"Reply to Notice - {matter_name}")
            st.download_button("📄 Download Reply as Word (.docx)", data=docx_file, file_name=f"Draft_Reply_{matter_name}.docx")

# ==========================================
# TAB 6: CHAT WITH PDF
# ==========================================
with tab_chat:
    st.header("💬 Interactive Q&A with Judgments")
    try:
        chat_judgments = [row for row in sh.worksheet("Judgments").get_all_records() if row.get("PDF File IDs")]
        if chat_judgments:
            c_dict = {r['Case Name']: r for r in chat_judgments}
            selected_chat_j = st.selectbox("Select a Judgment to Chat with:", options=list(c_dict.keys()))
            
            user_question = st.text_input("Ask a question about this specific judgment:")
            if st.button("Ask AI", key="chat_btn"):
                if user_question:
                    with st.spinner("Fetching PDF from Cloud Storage and analyzing..."):
                        file_ids = str(c_dict[selected_chat_j]["PDF File IDs"]).split(",")
                        doc_text = ""
                        for fid in file_ids:
                            if fid.strip():
                                pdf_bytes = download_from_gcs(fid.strip())
                                if pdf_bytes:
                                    reader = PyPDF2.PdfReader(BytesIO(pdf_bytes))
                                    for page in reader.pages:
                                        doc_text += page.extract_text() + "\n"
                        
                        if doc_text:
                            prompt = f"Based ONLY on the following legal judgment text, answer this question: {user_question}\n\nJudgment Text:\n{doc_text[:35000]}"
                            answer, err = ask_ai(prompt)
                            if not err:
                                st.session_state.chat_history.append({"q": user_question, "a": answer})
                        else:
                            st.error("Could not extract text from the Cloud Storage files.")
            for chat in reversed(st.session_state.chat_history):
                st.markdown(f"**🧑‍⚖️ You:** {chat['q']}")
                st.info(f"**🤖 AI:** {chat['a']}")
                st.markdown("---")
    except Exception as e:
        st.warning("Error fetching data.")